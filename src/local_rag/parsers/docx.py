"""DOCX text extraction using python-docx."""

import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class DocxDocument:
    """Parsed representation of a DOCX file."""

    text: str
    headings: list[str] = field(default_factory=list)
    tables_text: str = ""


def parse_docx(path: Path) -> DocxDocument:
    """Extract text, headings, and table content from a DOCX file.

    Args:
        path: Path to the DOCX file.

    Returns:
        DocxDocument with extracted content.
    """
    if not path.exists():
        logger.error("DOCX file not found: %s", path)
        return DocxDocument(text="")

    try:
        doc = Document(str(path))
    except Exception as e:
        logger.error("Failed to open DOCX %s: %s", path, e)
        return DocxDocument(text="")

    paragraphs: list[str] = []
    headings: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style and para.style.name and para.style.name.startswith("Heading"):
            headings.append(text)

        paragraphs.append(text)

    # Extract table text
    table_parts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                table_parts.append(" | ".join(row_cells))

    tables_text = "\n".join(table_parts)
    body_text = "\n\n".join(paragraphs)

    # Combine body and table text
    full_text = body_text
    if tables_text:
        full_text = full_text + "\n\n" + tables_text if full_text else tables_text

    return DocxDocument(
        text=full_text,
        headings=headings,
        tables_text=tables_text,
    )
